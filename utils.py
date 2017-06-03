# -*- coding: utf-8 -*-
import datetime
import fnmatch
import os
import random
import re
import time
import xml.etree.cElementTree as et

import docx

import config


class QuestionType(enumerate):
    choice = 'choice'
    s_choice = 'single_choice'
    m_choice = 'multi_choice'
    blank = 'blank'
    yes_no = 'yes_no'


def find_all_files(path):
    files = []
    for root, dirnames, filenames in os.walk(path):
        for filename in fnmatch.filter(filenames, '*.docx'):
            files.append(os.path.join(root, filename))
    return files


def get_file_name(path):
    last_slash = path.rfind('/')
    last_dot = path.rfind('.')
    return unicodify(path[last_slash + 1:last_dot])


def unicodify(unicode_string, encoding='utf-8'):
    return unicode(unicode_string, encoding)


def extract_from_docx(docx_path):
    doc = docx.Document(docx_path)
    doc_name = get_file_name(docx_path)
    root = et.Element('chapter')
    questions = et.SubElement(root, 'questions')

    i = 0
    while i < len(doc.paragraphs):
        line = doc.paragraphs[i].text.strip()

        # a question always starts with a number
        if line and line[0].isdigit():
            question = et.SubElement(questions, 'question')
            question.set('category', doc_name)
            question.text = line[re.search('(\d)+(.)?(\s)?', line, re.UNICODE).end():].strip()
            i += 1
            # leave out blank lines
            while not doc.paragraphs[i].text.strip():
                i += 1
            # answer is formatted to be right beneath the question
            answer = doc.paragraphs[i].text.strip()
            question.set('answer', answer)

            j = i + 1
            # an option starts with a letter, possibly followed by a delimiter
            regex_option = unicodify('^(\s)*[A-Za-z]+(\s|[.]|[、]|[。]|[，])*')
            ascii_A = 65
            while j < len(doc.paragraphs):
                text = doc.paragraphs[j].text.strip()
                if text and re.search(regex_option, text):
                    cur = doc.paragraphs[j].text.strip()
                    option = et.SubElement(question, 'option')
                    option.text = '%s. %s' % (chr(ascii_A), cur[re.search(regex_option, cur).end():])
                    ascii_A += 1
                    j += 1
                else:
                    break

            # 不能仅通过答案判断题型，如填空题答案为A
            if not question.getchildren():
                question.set('type', QuestionType.blank)
            else:
                a = answer.decode('ascii')
                if str(a).lower() in ['yes', 'no']:
                    question.set('type', QuestionType.yes_no)
                elif len(a) == 1:
                    question.set('type', QuestionType.s_choice)
                else:
                    question.set('type', QuestionType.m_choice)
            i = j
            continue
        i += 1

    tree = et.ElementTree(root)
    tree.write('./xml/xml_sub_%s.xml' % doc_name, encoding='utf-8', xml_declaration=True, method='xml')

    return et.tostring(root, encoding='utf-8', method='xml')


def create_question_xml(path):
    all_files = find_all_files(path=path)
    if not all_files:
        print 'No file found. Exiting'
        exit(1)
    root = et.Element('root')
    for f in all_files:
        print f
        xml = extract_from_docx(f)
        parsed = et.fromstring(xml)
        root.extend(parsed)
    tree = et.ElementTree(root)
    tree.write(config.xml_db, encoding='utf-8', xml_declaration=True, method='xml')


def add_to_category(category_dictionary, node):
    c = node.get('category')
    if c in category_dictionary:
        category_dictionary[c] += 1
    else:
        category_dictionary[c] = 1


def create_exam(xml, single_choice=None, multi_choice=None, blanks=None, yes_no=None, save_path='./exams'):
    root = None
    try:
        load_questions = et.parse(xml)
        root = load_questions.getroot()
    except IOError, e:
        print 'Failed to load xml:'
        print e.message
        exit(1)

    total_count = 0
    if single_choice:
        total_count += single_choice
    if multi_choice:
        total_count += multi_choice
    if blanks:
        total_count += blanks
    if yes_no:
        total_count += yes_no

    document = docx.Document()
    answer_doc = docx.Document()

    # styles may not work for unicode characters
    style = document.styles['Normal']
    font = style.font
    font.name = 'Source Han Sans CN'

    document.add_heading(unicodify('金融测试题'), 0)
    info = document.add_paragraph(unicodify('本次测试共%s题。 ' % total_count))
    info.add_run('@%s' % datetime.datetime.now().date())

    answer_doc.add_heading(unicodify('答案'), 0)

    category_count = {}

    if single_choice:
        list_choice = []
        for node in root.findall('.//question[@type="%s"]' % QuestionType.s_choice):
            list_choice.append(node)
        list_selected = random.sample(list_choice, single_choice)
        document.add_paragraph(unicodify('单项选择题'), style='Heading 1')
        for q in list_selected:
            add_to_category(category_dictionary=category_count, node=q)
            document.add_paragraph(q.text, style='List Number 2')
            answer_doc.add_paragraph('%s -- %s' %(q.get('answer'), q.get('category')), style='List Number')
            for o in q.findall('.//option'):
                document.add_paragraph(o.text, style='List 3')

    if multi_choice:
        list_choice = []
        for node in root.findall('.//question[@type="%s"]' % QuestionType.m_choice):
            list_choice.append(node)
        list_selected = random.sample(list_choice, multi_choice)
        document.add_paragraph(unicodify('多项选择题'), style='Heading 1')
        for q in list_selected:
            add_to_category(category_dictionary=category_count, node=q)
            document.add_paragraph(q.text, style='List Number 2')
            answer_doc.add_paragraph('%s -- %s' % (q.get('answer'), q.get('category')), style='List Number')
            for o in q.findall('.//option'):
                document.add_paragraph(o.text, style='List 3')

    if blanks:
        list_blanks = []
        for node in root.findall('.//question[@type="%s"]' % QuestionType.blank):
            list_blanks.append(node)
        list_selected = random.sample(list_blanks, blanks)
        document.add_paragraph(unicodify('填空题'), style='Heading 1')
        for q in list_selected:
            add_to_category(category_dictionary=category_count, node=q)
            document.add_paragraph(q.text, style='List Number 2')
            answer_doc.add_paragraph('%s -- %s' % (q.get('answer'), q.get('category')), style='List Number')

    if yes_no:
        list_yn = []
        for node in root.findall('.//question[@type="%s"]' % QuestionType.yes_no):
            list_yn.append(node)
        list_selected = random.sample(list_yn, yes_no)
        document.add_paragraph(unicodify('判断题'), style='Heading 1')
        for q in list_selected:
            add_to_category(category_dictionary=category_count, node=q)
            document.add_paragraph(q.text, style='List Number 2')
            answer_doc.add_paragraph('%s -- %s' % (q.get('answer'), q.get('category')), style='List Number')

    answer_doc.add_paragraph(unicodify('题型统计'), style='Heading 3')
    for key, value in category_count.items():
        answer_doc.add_paragraph(u'%s: %s题' % (key, value))

    document.save('%s/exam_%s.docx' % (save_path, time.strftime('%Y_%m_%d_%H_%M_%S')))
    answer_doc.save('%s/exams_answers_%s.docx' % (save_path, time.strftime('%Y_%m_%d_%H_%M_%S')))
    print 'Exam created with %s questions @ %s' % (total_count, time.strftime('%Y_%m_%d_%H_%M_%S'))
